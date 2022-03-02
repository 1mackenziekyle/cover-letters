from RPA.Browser.Selenium import Selenium
import time
from numpy import dtype 
import pandas as pd
from bs4 import BeautifulSoup
import docx

lib = Selenium()

frontend_words = [
    'front',
    'react',
    'javascript',
    'web',

]

backend_words = [
    'back',
    'pipeline',
    'data',
    'python',
    'R'
]

ORGANIZATION = "org"
LOCATION = "loc"
DATE = "date"


c_l_paragraphs = [
    "{ORGANIZATION}",
    "{LOCATION}, Canada",
    "{DATE}",
    "Dear Hiring Manager,",
    "As an Engineering student and aspiring Full Stack Developer, the space of developing software is extremely appealing and relatable to me. Moreover, creating engaging user-facing applications in new, up-to-date technologies including React is captivating for me and aligns with my job interests. Furthermore, I have a wide experience in many software technologies including Node.js, Python, and MATLAB. Furthermore, I am happy and eager to learn as much as possible and get familiar with the tools used at {JOB_TITLE}.",
    "While working at ICBC as a Full-Stack Developer Co-op, I sharpened my front-end design skills by developing a React app for a demo for a company website. Additionally, I worked on automation projects in Python using the Robot automation and testing Framework to optimize business processes and create internal solutions to expediate work being done by Insurance Broker. Throughout this co-op, I have learnt about working in a development team, and communicating technical ideas to different audiences and stakeholders. Additionally, this myriad of experiences gave me confidence in building production-level code.",
    "Additionally, as an Embedded Systems Developer on the UBC Solar Design Team, I developed low-level C code and gained an appreciation for analyzing algorithm complexity and creating a responsive system for real-time usage in a high stakes context. We used tools such as GitHub for version control in a large team environment. This introduction to data structures forged an interest inside me of optimization of algorithms.",
    "Finally, I gained front-end development experience through working at Charlene’s Web, a web development company. There, I worked alongside a WordPress developer to build a slideshow themed plugin for content creators to use as a tool to build their website. I developed the front end of the slideshow using Bootstrap 5, Express.js, and JQuery. and used industry-standard file transfer protocol to deploy the project online.",
    "The skills I have developed over years of programming in different environments have given me the tools I need to thrive as a developer at {ORGANIZATION}. I am confident that my wide array of experience in different topics would allow me to settle quickly into the team. I hope to bring vigor, a fresh outlook, and craftsmanship to work in my role. I look forward to hearing back from you and discussing the role further.",
    "Sincerely,",
    "Kyle Mackenzie"
]

def main(ubc_user, ubc_pw):
    try:
        # navigate
        navigate_to_browser(ubc_user=ubc_user, ubc_pw=ubc_pw)
        # get data
        soup = BeautifulSoup(lib.driver.page_source, "html.parser")
        table = pd.read_html(lib.driver.page_source)[0]
        # modify data
        lib.close_all_browsers()
        table.insert(1, "Job Type", "")
        # print to excel
        print(table.to_excel(excel_writer="output/table.xlsx", sheet_name="Jobs"))

        # generate documents
        generate_documents(table) 
    except Exception as e:
        print(e)
    finally: 
        print("=========DONE MAIN METHOD===========")





# Function: Open and navigate to SCOPE
def navigate_to_browser(ubc_user, ubc_pw):
    # OPEN BROWSER
    lib.open_available_browser("https://scope.sciencecoop.ubc.ca/students/cwl-current-student-login.htm")
    lib.set_browser_implicit_wait(10)

    # SCOPE LOGIN
    login_button = lib.click_link("xpath:/html/body/div/main/div/div/a")
    lib.wait_until_element_is_visible('xpath:/html/body/div[1]/div/div/div/div/h1')

    # CWL LOGIN
    username_input_field = "xpath://*[@id='username']"
    password_input_field = "xpath://*[@id='password']"
    pause(0.1)
    lib.input_text_when_element_is_visible(username_input_field, ubc_user)
    pause(0.1)
    lib.input_text_when_element_is_visible(password_input_field, ubc_pw)
    lib.click_button('xpath://*[@id="col2"]/form/div[3]/button')

    # SCOPE HOMEPAGE
    pause(0.5)
    lib.click_button_when_visible('xpath:/html/body/div[2]/header/div[4]/div/div/button')
    pause(1)
    lib.click_button_when_visible('xpath:/html/body/div[2]/header/div[3]/div[1]/nav/ul/li[2]/button')
    pause(1)
    lib.click_element_when_visible('xpath://*[@id="myAccountNav"]/nav/ul/li[2]/ul/li[2]/a')
    pause(3)
    lib.click_element_when_visible('xpath://*[@id="quickSearchCountsContainer"]/table[1]/tbody/tr[2]/td[2]/a')
    pause(1)






# Function: Generate Cover Letters
def generate_documents(table):
    # doc = docx.Document("cover-letter-template.docx")
    # para_count = 0
    # for para in doc.paragraphs:
    #     print("======paragraph #", para_count, "======")
    #     print(para.text)
    #     para_count += 1


    print("TEST: Writing to doc.")
    output_doc = docx.Document("input/header.docx")
    for paragraph in c_l_paragraphs:
        print("adding paragraph: ", paragraph, "\n")
        add_para = output_doc.add_paragraph(paragraph)
        print(add_para)









def pause(secs):
    time.sleep(secs)
    























































if __name__ == "__main__":
    main()
